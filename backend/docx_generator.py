"""DOCX generation utilities for meeting minutes."""

import os
import io
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from flask import send_file
from typing import Dict, Any, BinaryIO
from werkzeug.utils import secure_filename

# Configure logging
logger = logging.getLogger(__name__)

def generate_docx(data: Dict[str, Any]) -> BinaryIO:
    """
    Generate a DOCX document from meeting minutes data.
    
    Args:
        data: Meeting data dictionary with title, duration, summary, action_points, and transcription
        
    Returns:
        Flask send_file response with the generated DOCX
    """
    try:
        # Create a new document
        doc = Document()
        
        # Set up document styles
        _setup_document_styles(doc)
        
        # Add title with formatting
        title = doc.add_heading(data['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section with date and duration
        metadata_para = doc.add_paragraph(style='Metadata')
        metadata_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}\n")
        metadata_para.add_run(f"Duration: {data.get('duration', 'Not specified')}")
        
        # Add divider
        doc.add_paragraph('_' * 50, style='Subtle')
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(style='Normal')
        summary_para.add_run(data.get('summary', 'No summary available'))
        
        # Add action points with proper formatting
        action_points = data.get('action_points', [])
        if action_points and len(action_points) > 0:
            doc.add_heading('Action Points', level=1)
            
            for point in action_points:
                if point and point.strip():
                    # Use a proper bullet style
                    point_para = doc.add_paragraph(style='List Bullet')
                    point_para.add_run(point.strip())
                    point_para.paragraph_format.left_indent = Inches(0.25)
                    point_para.paragraph_format.first_line_indent = Inches(-0.25)
        else:
            doc.add_heading('Action Points', level=1)
            doc.add_paragraph('No action points recorded', style='Normal')
        
        # Add transcript with better formatting
        if data.get('transcription'):
            doc.add_heading('Transcript', level=1)
            transcript = data.get('transcription', '')
            
            # Split transcript into paragraphs for better readability
            paragraphs = transcript.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(style='Transcript')
                    p.add_run(para.strip())
        
        # Save document to memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Document name with date
        doc_name = secure_filename(data.get('title', 'meeting_minutes'))
        download_name = f"{doc_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        logger.info(f"Generated DOCX document for download: {download_name}")
        
        # Return file for download
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=download_name
        )
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        raise

def auto_save_docx(data: Dict[str, Any], original_filename: str = None, job_id: str = None) -> str:
    """
    Save DOCX to the docx_output_files directory.
    
    Args:
        data: Dictionary with meeting data
        original_filename: Original audio/video filename
        job_id: Unique identifier for the job (optional)
        
    Returns:
        Path to the saved DOCX file
    """
    try:
        # Create output directory if it doesn't exist
        docx_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'docx_output_files'))
        os.makedirs(docx_dir, exist_ok=True)
        
        # Create filename based on inputs
        if original_filename:
            base_name = os.path.splitext(os.path.basename(original_filename))[0]
            safe_name = secure_filename(base_name)
        else:
            safe_name = secure_filename(data.get("title", "meeting_minutes"))
        
        # Add job ID if available
        if job_id:
            docx_filename = f"{safe_name}_{job_id}.docx"
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            docx_filename = f"{safe_name}_{timestamp}.docx"
        
        filepath = os.path.join(docx_dir, docx_filename)
        
        # Handle duplicate filenames
        counter = 1
        while os.path.exists(filepath):
            name_base = f"{safe_name}_{counter}"
            if job_id:
                docx_filename = f"{name_base}_{job_id}.docx"
            else:
                docx_filename = f"{name_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(docx_dir, docx_filename)
            counter += 1
        
        # Create document
        doc = Document()
        
        # Set up document styles
        _setup_document_styles(doc)
        
        # Add title as document heading
        title = doc.add_heading(data.get('title', 'Meeting Minutes'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section with date and duration
        metadata_para = doc.add_paragraph(style='Metadata')
        metadata_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}\n")
        metadata_para.add_run(f"Duration: {data.get('duration', 'Not specified')}")
        
        # Add divider
        doc.add_paragraph('_' * 50, style='Subtle')
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(style='Normal')
        summary_para.add_run(data.get('summary', 'No summary available'))
        
        # Add action points section with proper formatting
        action_points = data.get('action_points', [])
        doc.add_heading('Action Points', level=1)
        
        if action_points and len(action_points) > 0:
            for point in action_points:
                if point and point.strip():
                    # Use proper bullet formatting
                    point_para = doc.add_paragraph(style='List Bullet')
                    point_para.add_run(point.strip())
                    point_para.paragraph_format.left_indent = Inches(0.25)
                    point_para.paragraph_format.first_line_indent = Inches(-0.25)
        else:
            doc.add_paragraph('No action points recorded', style='Normal')
        
        # Add transcript section with better formatting
        if data.get('transcription'):
            doc.add_heading('Transcript', level=1)
            transcript = data.get('transcription', '')
            
            # Split transcript into paragraphs for better readability
            paragraphs = transcript.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(style='Transcript')
                    p.add_run(para.strip())
        
        # Save to file
        doc.save(filepath)
        
        logger.info(f"DOCX saved to: {filepath}")
        return filepath
        
    except Exception as e:
        logger.error(f"Error saving DOCX: {str(e)}")
        raise

def _setup_document_styles(doc: Document) -> None:
    """Set up document styles for consistent formatting"""
    # Title style
    title_style = doc.styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
    
    # Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(10)
    
    # Bullet list style
    if 'List Bullet' not in doc.styles:
        list_style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = doc.styles['Normal']
    else:
        list_style = doc.styles['List Bullet']
    
    list_style.font.size = Pt(11)
    
    # Create metadata style
    metadata_style = doc.styles.add_style('Metadata', WD_STYLE_TYPE.PARAGRAPH)
    metadata_style.font.size = Pt(10)
    metadata_style.font.italic = True
    metadata_style.paragraph_format.space_after = Pt(8)
    
    # Subtle style for dividers
    subtle_style = doc.styles.add_style('Subtle', WD_STYLE_TYPE.PARAGRAPH)
    subtle_style.font.size = Pt(8)
    subtle_style.font.color.rgb = RGBColor(150, 150, 150)
    subtle_style.paragraph_format.space_before = Pt(5)
    subtle_style.paragraph_format.space_after = Pt(15)
    
    # Transcript style
    transcript_style = doc.styles.add_style('Transcript', WD_STYLE_TYPE.PARAGRAPH)
    transcript_style.base_style = doc.styles['Normal']
    transcript_style.font.size = Pt(10)
    transcript_style.paragraph_format.space_after = Pt(6)
    transcript_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE